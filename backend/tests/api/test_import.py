"""M1-T2 acceptance: file import (FR-1.1, API-4)."""

import pytest
from httpx import AsyncClient


class TestFileImport:
    """API-4: POST /projects/{id}/source/import — TXT/DOCX file import."""

    @pytest.fixture
    async def project_id(self, app_client: AsyncClient) -> str:
        resp = await app_client.post(
            "/api/v1/projects",
            json={"title": "import-test"},
        )
        return resp.json()["id"]

    # -- TXT import ------------------------------------------------------------

    async def test_import_txt_utf8(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """Import a UTF-8 TXT file and get chapter preview."""
        content = "第一章 旧书店\n\n林晚推开旧书店的门。\n\n灰尘在夕阳中扬起。"
        resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import",
            files={"file": ("novel.txt", content.encode("utf-8"), "text/plain")},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "chapters" in data
        assert len(data["chapters"]) >= 1

    async def test_import_txt_with_chapter_markers(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """TXT with 第X章 markers auto-detects boundaries."""
        content = (
            "第一章 初遇\n\n她走进了旧书店。\n\n"
            "第二章 重逢\n\n他在书架后面等她。\n\n"
            "第三章 真相\n\n一切谜底都揭开了。"
        )
        resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import",
            files={"file": ("novel.txt", content.encode("utf-8"), "text/plain")},
        )
        assert resp.status_code == 200
        data = resp.json()
        # Should detect 3 chapters
        assert len(data["chapters"]) == 3
        assert data["chapters"][0]["text"].startswith("第一章 初遇")

    async def test_import_txt_gbk_encoding(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """Import a TXT file with GBK encoding."""
        content = "第一章 开端\n\n故事从这里开始。"
        resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import",
            files={
                "file": (
                    "novel.txt",
                    content.encode("gbk"),
                    "text/plain",
                )
            },
        )
        assert resp.status_code == 200
        data = resp.json()
        assert len(data["chapters"]) >= 1

    async def test_import_txt_no_chapter_markers(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """TXT without chapter markers returns single chapter."""
        content = "只是一段简单的文字。\n\n没有任何章节标记。"
        resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import",
            files={"file": ("novel.txt", content.encode("utf-8"), "text/plain")},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert len(data["chapters"]) == 1

    async def test_import_empty_file_rejected(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """Empty file returns 400."""
        resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import",
            files={"file": ("empty.txt", b"   \n  ", "text/plain")},
        )
        assert resp.status_code == 400

    async def test_import_unsupported_format(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """Unsupported format returns 400."""
        resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import",
            files={"file": ("novel.pdf", b"dummy", "application/pdf")},
        )
        assert resp.status_code == 400

    async def test_import_docx(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """Import a minimal DOCX file."""
        # Build a minimal DOCX in memory
        from io import BytesIO

        from docx import Document

        doc = Document()
        doc.add_paragraph("第一章 旧书店")
        doc.add_paragraph("林晚推开了门。")
        doc.add_paragraph("她站了很久。")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import",
            files={
                "file": (
                    "novel.docx",
                    buf.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert resp.status_code == 200
        data = resp.json()
        assert len(data["chapters"]) >= 1

    async def test_confirm_txt_import_persists_chapters(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """Confirming import preview persists chapters into source view."""
        content = (
            "第一章 初遇\n\n她走进了旧书店。\n\n"
            "第二章 重逢\n\n他在书架后面等她。\n\n"
            "第三章 真相\n\n一切谜底都揭开了。"
        )
        import_resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import",
            files={"file": ("novel.txt", content.encode("utf-8"), "text/plain")},
        )
        assert import_resp.status_code == 200

        confirm_resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import:confirm",
            json={"chapters": import_resp.json()["chapters"]},
        )

        assert confirm_resp.status_code == 200
        confirmed = confirm_resp.json()
        assert len(confirmed["chapters"]) == 3
        assert confirmed["threshold"]["passed"] is True
        assert confirmed["chapters"][0]["paragraphs"][0]["text"] == "第一章 初遇"
        assert confirmed["chapters"][0]["paragraphs"][1]["text"] == "她走进了旧书店。"

        source_resp = await app_client.get(f"/api/v1/projects/{project_id}/source")
        assert source_resp.json()["stats"]["chapter_count"] == 3

    async def test_confirm_import_replaces_existing_source(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """Confirmed import replaces any existing source paragraphs."""
        await app_client.post(
            f"/api/v1/projects/{project_id}/source/chapters",
            json={"title": "旧章", "text": "旧内容。", "order": 1},
        )

        resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import:confirm",
            json={
                "chapters": [
                    {"title": "新章", "text": "新内容。", "order": 1},
                ]
            },
        )

        assert resp.status_code == 200
        data = resp.json()
        assert len(data["chapters"]) == 1
        assert data["chapters"][0]["paragraphs"][0]["text"] == "新内容。"

    async def test_confirm_docx_import_persists_source(
        self, app_client: AsyncClient, project_id: str
    ) -> None:
        """DOCX import preview can be confirmed into source view."""
        from io import BytesIO

        from docx import Document

        doc = Document()
        doc.add_paragraph("第一章 旧书店")
        doc.add_paragraph("林晚推开了门。")
        doc.add_paragraph("第二章 归来")
        doc.add_paragraph("她再次站在门口。")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        import_resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import",
            files={
                "file": (
                    "novel.docx",
                    buf.read(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert import_resp.status_code == 200

        confirm_resp = await app_client.post(
            f"/api/v1/projects/{project_id}/source/import:confirm",
            json={"chapters": import_resp.json()["chapters"]},
        )

        assert confirm_resp.status_code == 200
        assert confirm_resp.json()["stats"]["chapter_count"] == 2

    async def test_project_not_found(
        self, app_client: AsyncClient
    ) -> None:
        """Non-existent project returns 404."""
        resp = await app_client.post(
            "/api/v1/projects/nonexistent/source/import",
            files={"file": ("novel.txt", b"content", "text/plain")},
        )
        assert resp.status_code == 404
