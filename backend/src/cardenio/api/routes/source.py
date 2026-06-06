"""Source (novel import) API (api.md §4, API-3~6).

M1-T1: chapter input (paste/type) — done
M1-T2: file import (TXT / DOCX) — done
M1-T3: chapter segmentation — done
M1-T4: threshold check — done
M1-T5: text cleaning — current
"""

from __future__ import annotations

import io
import re

from fastapi import APIRouter, Depends, HTTPException, UploadFile
from fastapi.responses import JSONResponse

from cardenio.api.deps import get_artifact_store
from cardenio.domain.models.source import (
    Chapter,
    ConfirmImportRequest,
    CreateChapterRequest,
    SourceParagraph,
    SourceStats,
)
from cardenio.storage.sqlite_store import SqliteArtifactStore

router = APIRouter(prefix="/projects/{project_id}/source")

ALLOWED_MIMES = {
    "text/plain": "txt",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}

_CHAPTER_PATTERNS = [
    r"第[一二三四五六七八九十百千零\d]+[章节]",
    r"chapter\s*\d+",
    r"第[一二三四五六七八九十百千零\d]+卷",
]
_CHAPTER_RE = re.compile("|".join(_CHAPTER_PATTERNS), re.IGNORECASE)


# =============================================================================
# T1 — Chapter CRUD
# =============================================================================


@router.post("/chapters", status_code=201)
async def create_chapter(
    project_id: str,
    body: CreateChapterRequest,
    store: SqliteArtifactStore = Depends(get_artifact_store),
) -> dict:
    """API-3: Add a chapter — persists text and builds paragraph index."""
    proj = await store.get_project(project_id)
    if proj is None:
        raise HTTPException(status_code=404, detail="Project not found")

    existing = await store.list_chapters(project_id)
    order = len(existing) + 1 if not body.order else body.order

    cleaned_text = _clean_basic(body.text)
    raw_paragraphs = _split_paragraphs(cleaned_text)
    chapter_id = f"ch_{order}"
    paragraph_models: list[SourceParagraph] = []
    for idx, para_text in enumerate(raw_paragraphs):
        paragraph_models.append(SourceParagraph(index=idx + 1, text=para_text))

    await store.save_paragraphs(
        project_id=project_id,
        chapter_id=chapter_id,
        paragraphs=[p.model_dump(mode="json") for p in paragraph_models],
    )

    return {
        "id": chapter_id,
        "title": body.title,
        "order": order,
        "char_count": _char_sum(paragraph_models),
        "paragraphs": [p.model_dump(mode="json") for p in paragraph_models],
    }


@router.get("")
async def get_source(
    project_id: str,
    store: SqliteArtifactStore = Depends(get_artifact_store),
) -> dict:
    """API-5: Get all chapters with paragraph index and threshold check.

    FR-1.3: threshold.blocked=true when < 3 chapters — downstream
    generation endpoints will return 409 chapter_threshold_unmet.
    """
    proj = await store.get_project(project_id)
    if proj is None:
        raise HTTPException(status_code=404, detail="Project not found")

    chapters = await store.list_chapters(project_id)
    total_chars = sum(c["char_count"] for c in chapters)
    stats = SourceStats(chapter_count=len(chapters), char_count=total_chars)

    return {
        "chapters": chapters,
        "stats": stats.model_dump(mode="json"),
        "threshold": {
            "min_chapters": 3,
            "passed": stats.threshold_passed,
            "blocked": not stats.threshold_passed,
        },
    }


@router.get("/threshold")
async def get_threshold(
    project_id: str,
    store: SqliteArtifactStore = Depends(get_artifact_store),
) -> dict:
    """Check whether the project meets the 3-chapter minimum (FR-1.3).

    Called before attempting any downstream generation action.
    Returns 200 with threshold status, or 409 if blocked.
    """
    chapters = await store.list_chapters(project_id)
    total_chars = sum(c["char_count"] for c in chapters)
    stats = SourceStats(chapter_count=len(chapters), char_count=total_chars)

    result = {
        "min_chapters": 3,
        "current_chapters": stats.chapter_count,
        "current_chars": stats.char_count,
        "passed": stats.threshold_passed,
    }

    if not stats.threshold_passed:
        return JSONResponse(
            status_code=409,
            content={
                "error": {
                    "code": "chapter_threshold_unmet",
                    "message": f"Need at least 3 chapters; found {stats.chapter_count}",
                    "retryable": False,
                    "details": result,
                }
            },
        )

    return result


# =============================================================================
# T2 — File import
# =============================================================================


@router.post("/import")
async def import_file(
    project_id: str,
    file: UploadFile,
    store: SqliteArtifactStore = Depends(get_artifact_store),
) -> dict:
    """API-4: Import TXT/DOCX file with auto-chapter detection.

    Returns auto-segmented chapter preview for author review.
    Chapters are NOT persisted until confirmed via :import:confirm (M1-T3).
    """
    proj = await store.get_project(project_id)
    if proj is None:
        raise HTTPException(status_code=404, detail="Project not found")

    if file.content_type is None:
        raise HTTPException(status_code=400, detail="Unknown file type")

    raw_bytes = await file.read()

    try:
        text, warnings = _extract_text(raw_bytes, file.content_type, file.filename or "")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e

    if not text.strip():
        raise HTTPException(status_code=400, detail="File contains no readable text")

    chapters = _detect_chapters(text)

    return {"chapters": chapters, "warnings": warnings}


@router.post("/import:confirm")
async def confirm_import(
    project_id: str,
    body: ConfirmImportRequest,
    store: SqliteArtifactStore = Depends(get_artifact_store),
) -> dict:
    """Persist an edited import preview as source chapters."""
    proj = await store.get_project(project_id)
    if proj is None:
        raise HTTPException(status_code=404, detail="Project not found")

    await store.delete_all_paragraphs(project_id)

    for idx, chapter in enumerate(body.chapters):
        order = chapter.order or idx + 1
        chapter_id = f"ch_{order}"
        paragraphs = [
            SourceParagraph(index=p_idx + 1, text=paragraph)
            for p_idx, paragraph in enumerate(_split_paragraphs(chapter.text))
        ]
        if not paragraphs:
            continue
        await store.save_paragraphs(
            project_id=project_id,
            chapter_id=chapter_id,
            paragraphs=[p.model_dump(mode="json") for p in paragraphs],
        )

    return await get_source(project_id, store)


# =============================================================================
# T3 — Chapter segmentation (edit / delete / merge / split)
# =============================================================================


@router.put("/chapters/{chapter_id}")
async def update_chapter(
    project_id: str,
    chapter_id: str,
    chapter: Chapter,
    store: SqliteArtifactStore = Depends(get_artifact_store),
) -> dict:
    """API-6: Edit a chapter. Replaces all paragraphs for the chapter."""
    proj = await store.get_project(project_id)
    if proj is None:
        raise HTTPException(status_code=404, detail="Project not found")

    # Delete old paragraphs and re-insert
    await store.delete_paragraphs(project_id, chapter_id)

    paragraph_models: list[SourceParagraph] = []
    for para in chapter.paragraphs:
        paragraph_models.append(
            SourceParagraph(index=para.index, text=_clean_basic(para.text))
        )

    await store.save_paragraphs(
        project_id=project_id,
        chapter_id=chapter_id,
        paragraphs=[p.model_dump(mode="json") for p in paragraph_models],
    )

    return {
        "id": chapter_id,
        "title": chapter.title,
        "order": chapter.order,
        "char_count": chapter.char_count,
        "paragraphs": [p.model_dump(mode="json") for p in paragraph_models],
    }


@router.delete("/chapters/{chapter_id}", status_code=204)
async def delete_chapter(
    project_id: str,
    chapter_id: str,
    store: SqliteArtifactStore = Depends(get_artifact_store),
) -> None:
    """API-6: Delete a chapter and its paragraphs."""
    proj = await store.get_project(project_id)
    if proj is None:
        raise HTTPException(status_code=404, detail="Project not found")

    deleted = await store.delete_paragraphs(project_id, chapter_id)
    if deleted == 0:
        raise HTTPException(status_code=404, detail="Chapter not found")


@router.post("/chapters:resegment")
async def resegment_chapters(
    project_id: str,
    body: dict,
    store: SqliteArtifactStore = Depends(get_artifact_store),
) -> dict:
    """API-6: Split or merge chapters.

    Accepts:
        {"op": "split",  "chapter_id": "ch_2", "at_paragraph": 47}
        {"op": "merge",  "chapter_ids": ["ch_1", "ch_2"], "new_title": "合并章"}

    Paragraph indices are remapped after the operation.
    """
    proj = await store.get_project(project_id)
    if proj is None:
        raise HTTPException(status_code=404, detail="Project not found")

    op = body.get("op")
    if op not in ("split", "merge"):
        raise HTTPException(status_code=400, detail="op must be 'split' or 'merge'")

    if op == "split":
        await _split_chapter(store, project_id, body)
    elif op == "merge":
        await _merge_chapters(store, project_id, body)

    # Return updated chapters
    return await get_source(project_id, store)


async def _split_chapter(
    store: SqliteArtifactStore, project_id: str, body: dict
) -> None:
    """Split one chapter into two at the given paragraph boundary."""
    chapter_id: str = body["chapter_id"]
    at_paragraph: int = body["at_paragraph"]

    rows = await store.get_paragraphs(project_id, chapter_id=chapter_id)
    if not rows:
        raise HTTPException(status_code=404, detail="Chapter not found")

    part1 = [r for r in rows if r["paragraph_index"] < at_paragraph]
    part2 = [r for r in rows if r["paragraph_index"] >= at_paragraph]

    if not part2:
        raise HTTPException(
                status_code=400, detail="Split point is chapter boundary — nothing to split"
            )

    # Delete old chapter
    await store.delete_paragraphs(project_id, chapter_id)

    # Write part 1 with original chapter_id
    await store.save_paragraphs(
        project_id=project_id,
        chapter_id=chapter_id,
        paragraphs=[
            {"index": idx + 1, "text": p["text"]} for idx, p in enumerate(part1)
        ],
    )

    # Write part 2 as a new chapter
    chapters = await store.list_chapters(project_id)
    new_order = len(chapters) + 2  # one deleted, plus new one coming
    new_id = f"ch_{new_order}"
    await store.save_paragraphs(
        project_id=project_id,
        chapter_id=new_id,
        paragraphs=[
            {"index": idx + 1, "text": p["text"]} for idx, p in enumerate(part2)
        ],
    )


async def _merge_chapters(
    store: SqliteArtifactStore, project_id: str, body: dict
) -> None:
    """Merge two or more chapters into one."""
    chapter_ids: list[str] = body["chapter_ids"]

    if len(chapter_ids) < 2:
        raise HTTPException(status_code=400, detail="Need at least 2 chapters to merge")

    all_paragraphs: list[dict] = []
    for cid in chapter_ids:
        rows = await store.get_paragraphs(project_id, chapter_id=cid)
        all_paragraphs.extend(rows)
        await store.delete_paragraphs(project_id, cid)

    # Write merged chapter
    merged_id = chapter_ids[0]
    await store.save_paragraphs(
        project_id=project_id,
        chapter_id=merged_id,
        paragraphs=[
            {"index": idx + 1, "text": p["text"]}
            for idx, p in enumerate(all_paragraphs)
        ],
    )


# =============================================================================
# Text extraction
# =============================================================================


def _extract_text(raw: bytes, mime: str, filename: str) -> tuple[str, list[str]]:
    warnings: list[str] = []
    ext = ALLOWED_MIMES.get(mime) or _ext_from_filename(filename)

    if ext == "txt":
        text = _extract_txt(raw)
    elif ext == "docx":
        text = _extract_docx(raw)
    else:
        supported = ", ".join(ALLOWED_MIMES.values())
        raise ValueError(f"Unsupported format: {ext or mime}. Supported: {supported}")

    text = _clean_basic(text)
    return text, warnings


def _ext_from_filename(filename: str) -> str:
    if filename.lower().endswith(".txt"):
        return "txt"
    if filename.lower().endswith(".docx"):
        return "docx"
    return ""


def _extract_txt(raw: bytes) -> str:
    for enc in ("utf-8", "gbk", "gb2312", "gb18030"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_docx(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx is required for DOCX import") from None

    doc = Document(io.BytesIO(raw))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    return "\n\n".join(paragraphs)


# =============================================================================
# Chapter detection + text helpers
# =============================================================================


def _detect_chapters(text: str) -> list[dict]:
    """Auto-detect chapter boundaries from raw text."""
    lines = text.split("\n")
    chapter_starts: list[int] = []

    for i, line in enumerate(lines):
        stripped = line.strip()
        if _CHAPTER_RE.search(stripped) and len(stripped) < 80:
            chapter_starts.append(i)

    if not chapter_starts:
        total_chars = sum(len(ln.strip()) for ln in lines if ln.strip())
        paragraphs = _split_paragraphs(text)
        return [{
            "title": "第一章",
            "char_count": total_chars,
            "paragraphs": [1, max(1, len(paragraphs))],
            "text": text,
        }]

    chapters: list[dict] = []
    for idx, start_line in enumerate(chapter_starts):
        end_line = (
            chapter_starts[idx + 1] if idx + 1 < len(chapter_starts) else len(lines)
        )
        chapter_text = "\n".join(lines[start_line:end_line])
        paras = _split_paragraphs(chapter_text)
        title = lines[start_line].strip()

        chapters.append({
            "title": title,
            "char_count": sum(len(p) for p in paras),
            "paragraphs": [1, max(1, len(paras))],
            "text": chapter_text,
        })

    return [c for c in chapters if c["char_count"] > 0]


# =============================================================================
# M1-T5 — text cleaning (FR-1.4)
# =============================================================================


def _clean_basic(text: str) -> str:
    """Basic text cleaning.  M1-T5: preserves intentional whitespace.

    Steps (order matters):
    1. BOM removal
    2. CRLF / CR → LF
    3. Remove control chars except \n, \t
    4. Collapse 3+ consecutive newlines to 2 (keep paragraph breaks)
    5. Full-width → half-width for ASCII variants
    6. Strip trailing whitespace per line (keep leading indentation for prose)
    """
    # 1. Remove BOM
    text = text.lstrip("\ufeff")

    # 2. Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 3. Remove control characters except \n (0x0A) and \t (0x09)
    text = "".join(
        ch for ch in text
        if ch == "\n" or ch == "\t" or ord(ch) >= 0x20 or ch == "　"
    )

    # 4. Collapse 3+ consecutive newlines to 2 (preserve paragraph gaps)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 5. Full-width → half-width for ASCII variants. Non-ASCII Chinese
    # punctuation such as "。" and "《》" stays unchanged.
    text = _normalize_fullwidth_ascii(text)

    # 6. Strip trailing whitespace per line (preserve leading indent for prose)
    lines = text.split("\n")
    text = "\n".join(line.rstrip() for line in lines)

    return text


def _normalize_fullwidth_ascii(text: str) -> str:
    """Normalize full-width ASCII variants to their half-width forms."""
    chinese_punctuation = {
        "，",
        "。",
        "！",
        "？",
        "；",
        "：",
        "、",
        "（",
        "）",
        "《",
        "》",
        "〈",
        "〉",
        "【",
        "】",
        "“",
        "”",
        "‘",
        "’",
    }
    normalized: list[str] = []
    for ch in text:
        code = ord(ch)
        if ch == "　":
            normalized.append(" ")
        elif ch in chinese_punctuation:
            normalized.append(ch)
        elif 0xFF01 <= code <= 0xFF5E:
            normalized.append(chr(code - 0xFEE0))
        else:
            normalized.append(ch)
    return "".join(normalized)


def _split_paragraphs(text: str) -> list[str]:
    blocks = text.strip().split("\n\n")
    return [b.strip() for b in blocks if b.strip()]


def _char_sum(paras: list[SourceParagraph]) -> int:
    return sum(len(p.text) for p in paras)
